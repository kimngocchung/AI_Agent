"""
Document Processor
Process uploaded documents and add to FAISS index
"""

import os
from typing import List
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredMarkdownLoader
)
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
import tempfile

# === ĐƯỜNG DẪN CHUẨN ===
def get_faiss_index_path():
    """Trả về đường dẫn tuyệt đối tới FAISS index - GIỐNG với retriever.py"""
    # Từ utils/document_processor.py -> lên 1 cấp về thư mục gốc
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_dir, 'my_faiss_index')

# Đường dẫn mặc định
DEFAULT_FAISS_PATH = get_faiss_index_path()


def process_pdf(file_bytes, filename: str) -> List[Document]:
    """Extract text from PDF file"""
    try:
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        loader = PyPDFLoader(tmp_path)
        docs = loader.load()
        
        # Clean up
        os.unlink(tmp_path)
        
        # Add metadata
        for doc in docs:
            doc.metadata['source'] = filename
            
        return docs
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return []

def process_txt(file_bytes, filename: str) -> List[Document]:
    """Read TXT file"""
    try:
        text = file_bytes.decode('utf-8')
        doc = Document(page_content=text, metadata={'source': filename})
        return [doc]
    except Exception as e:
        print(f"Error processing TXT: {e}")
        return []

def process_docx(file_bytes, filename: str) -> List[Document]:
    """Extract text from DOCX file"""
    try:
        from docx import Document as DocxDocument
        
        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        doc = DocxDocument(tmp_path)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        
        # Clean up
        os.unlink(tmp_path)
        
        return [Document(page_content=text, metadata={'source': filename})]
    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return []

def process_md(file_bytes, filename: str) -> List[Document]:
    """Read Markdown file"""
    try:
        text = file_bytes.decode('utf-8')
        doc = Document(page_content=text, metadata={'source': filename})
        return [doc]
    except Exception as e:
        print(f"Error processing MD: {e}")
        return []

def chunk_documents(docs: List[Document], chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """Split documents into chunks"""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    return text_splitter.split_documents(docs)

def add_to_faiss(docs: List[Document], faiss_path: str = None) -> tuple[bool, str]:
    """
    Add documents to FAISS index
    
    Args:
        docs: List of documents to add
        faiss_path: Path to FAISS index directory (mặc định dùng đường dẫn chuẩn)
        
    Returns:
        (success: bool, message: str)
    """
    # Sử dụng đường dẫn chuẩn nếu không chỉ định
    if faiss_path is None or faiss_path == "my_faiss_index":
        faiss_path = DEFAULT_FAISS_PATH
    
    print(f"--- [FAISS] Adding {len(docs)} docs to: {faiss_path} ---")
    
    try:
        # Initialize embeddings
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        # Check if index exists
        if os.path.exists(faiss_path):
            # Load existing index
            vectorstore = FAISS.load_local(
                faiss_path,
                embeddings,
                allow_dangerous_deserialization=True
            )
            
            old_count = vectorstore.index.ntotal
            
            # Add new documents
            vectorstore.add_documents(docs)
            
            # Save updated index
            vectorstore.save_local(faiss_path)
            
            new_count = vectorstore.index.ntotal
            print(f"--- [FAISS] Updated: {old_count} -> {new_count} docs ---")
            
            # === RELOAD RETRIEVER ===
            try:
                # Import từ utils/ -> cần dùng đường dẫn tuyệt đối
                import sys
                import importlib
                # Thêm thư mục gốc vào path nếu chưa có
                base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                if base_dir not in sys.path:
                    sys.path.insert(0, base_dir)
                
                # Import và reload
                retriever_module = importlib.import_module('core.chains.retriever')
                retriever_module.reload_retriever()
                print("--- [FAISS] Retriever reloaded! ---")
            except Exception as e:
                print(f"--- [FAISS Warning] Could not reload retriever: {e} ---")
            
            return True, f"✅ Added {len(docs)} chunks to existing FAISS index (total: {new_count})"
        else:
            # Create directory if not exists
            os.makedirs(os.path.dirname(faiss_path), exist_ok=True)
            
            # Create new index
            vectorstore = FAISS.from_documents(docs, embeddings)
            vectorstore.save_local(faiss_path)
            
            print(f"--- [FAISS] Created new index with {len(docs)} docs ---")
            
            # === RELOAD RETRIEVER ===
            try:
                import sys
                import importlib
                base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                if base_dir not in sys.path:
                    sys.path.insert(0, base_dir)
                retriever_module = importlib.import_module('core.chains.retriever')
                retriever_module.reload_retriever()
                print("--- [FAISS] Retriever reloaded! ---")
            except Exception as e:
                print(f"--- [FAISS Warning] Could not reload retriever: {e} ---")
            
            return True, f"✅ Created new FAISS index with {len(docs)} chunks"
            
    except Exception as e:
        import traceback
        traceback.print_exc()
        return False, f"❌ Error adding to FAISS: {str(e)}"

def get_faiss_stats(faiss_path: str = None) -> dict:
    """
    Get statistics about FAISS index
    
    Returns:
        dict with stats (exists, doc_count, etc.)
    """
    # Sử dụng đường dẫn chuẩn nếu không chỉ định
    if faiss_path is None or faiss_path == "my_faiss_index":
        faiss_path = DEFAULT_FAISS_PATH
    
    try:
        if not os.path.exists(faiss_path):
            return {"exists": False, "doc_count": 0, "path": faiss_path}
        
        embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        
        vectorstore = FAISS.load_local(
            faiss_path,
            embeddings,
            allow_dangerous_deserialization=True
        )
        
        # Get document count
        doc_count = vectorstore.index.ntotal
        
        return {
            "exists": True,
            "doc_count": doc_count,
            "path": faiss_path
        }
    except Exception as e:
        print(f"Error getting FAISS stats: {e}")
        return {"exists": False, "doc_count": 0, "error": str(e), "path": faiss_path}